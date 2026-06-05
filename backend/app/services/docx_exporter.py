import re
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


HEADING_RE = re.compile(r"^(#{1,6})\s+(.+)$")
PUBLIC_ACCOUNT_TITLE_RE = re.compile(r"^【[^】]+】.+$")
ORDERED_LIST_RE = re.compile(r"^\d+[\.)]\s+(.+)$")
BULLET_LIST_RE = re.compile(r"^[-*+]\s+(.+)$")
INLINE_TOKEN_RE = re.compile(r"(\$\$[^$]+\$\$|\$[^$]+\$|\*\*[^*]+\*\*|__[^_]+__|`[^`]+`|\*[^*]+\*|\[[^\]]+\]\([^)]+\))")
LINK_RE = re.compile(r"^\[([^\]]+)\]\(([^)]+)\)$")
LATEX_SYMBOLS = {
    r"\approx": "≈",
    r"\sim": "∼",
    r"\leq": "≤",
    r"\geq": "≥",
    r"\neq": "≠",
    r"\times": "×",
    r"\cdot": "·",
    r"\alpha": "α",
    r"\beta": "β",
    r"\gamma": "γ",
    r"\delta": "δ",
    r"\Delta": "Δ",
    r"\lambda": "λ",
    r"\mu": "μ",
    r"\rho": "ρ",
    r"\sigma": "σ",
    r"\sum": "Σ",
}
SUPERSCRIPT_DIGITS = str.maketrans("0123456789+-=()", "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾")
DEFAULT_DOCUMENT_TITLE = "研读非洲｜第X期"
WORD_FONT_NAME = "SimSun"


class DocxExporter:
    def export_markdown(self, markdown: str, output_path: Path, default_title: str = DEFAULT_DOCUMENT_TITLE) -> Path:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        document = Document()
        self._configure_document_styles(document)
        in_formula_block = False
        title_written = False
        for raw_line in markdown.splitlines():
            line = raw_line.strip()
            if not line:
                continue
            if not title_written:
                title = self._title_from_first_line(line)
                if title is not None:
                    self._add_title(document, title)
                    title_written = True
                    continue
                self._add_title(document, default_title)
                title_written = True
            if line == "$$":
                in_formula_block = not in_formula_block
                continue
            if in_formula_block:
                paragraph = document.add_paragraph()
                self._add_run(paragraph, self._readable_latex(line))
                continue
            if set(line) <= {"-", "*", "_"} and len(line) >= 3:
                continue
            heading = HEADING_RE.match(line)
            ordered = ORDERED_LIST_RE.match(line)
            bullet = BULLET_LIST_RE.match(line)
            if heading:
                level = min(len(heading.group(1)), 6)
                paragraph = document.add_heading("", level=level)
                self._add_inline_runs(paragraph, heading.group(2).strip())
            elif ordered:
                paragraph = document.add_paragraph(style="List Number")
                self._add_inline_runs(paragraph, ordered.group(1).strip())
            elif bullet:
                paragraph = document.add_paragraph(style="List Bullet")
                self._add_inline_runs(paragraph, bullet.group(1).strip())
            elif line.startswith("> "):
                paragraph = document.add_paragraph(style="Intense Quote")
                self._add_inline_runs(paragraph, line[2:].strip())
            else:
                paragraph = document.add_paragraph()
                self._add_inline_runs(paragraph, line)
        if not title_written:
            self._add_title(document, default_title)
        document.save(output_path)
        return output_path

    def _configure_document_styles(self, document: Document) -> None:
        self._set_style_font(document.styles["Normal"], WORD_FONT_NAME, size=12)
        self._set_style_font(document.styles["Title"], WORD_FONT_NAME, size=20, bold=True)
        for level in range(1, 7):
            self._set_style_font(document.styles[f"Heading {level}"], WORD_FONT_NAME, bold=True)
        for style_name in ["List Number", "List Bullet", "Intense Quote"]:
            if style_name in document.styles:
                self._set_style_font(document.styles[style_name], WORD_FONT_NAME)

    def _set_style_font(self, style, font_name: str, size: int | None = None, bold: bool | None = None) -> None:
        style.font.name = font_name
        if size is not None:
            style.font.size = Pt(size)
        if bold is not None:
            style.font.bold = bold
        self._set_element_font(style.element, font_name)

    def _set_element_font(self, element, font_name: str) -> None:
        r_pr = element.get_or_add_rPr()
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for attribute in ["ascii", "hAnsi", "eastAsia", "cs"]:
            r_fonts.set(qn(f"w:{attribute}"), font_name)

    def _add_title(self, document: Document, title: str) -> None:
        paragraph = document.add_paragraph(style="Title")
        self._add_inline_runs(paragraph, title)
        document.core_properties.title = title

    def _title_from_first_line(self, line: str) -> str | None:
        heading = HEADING_RE.match(line)
        if heading and len(heading.group(1)) == 1:
            return heading.group(2).strip()
        if PUBLIC_ACCOUNT_TITLE_RE.match(line):
            return line
        return None

    def _add_inline_runs(self, paragraph, text: str) -> None:
        for token in INLINE_TOKEN_RE.split(text):
            if not token:
                continue
            link = LINK_RE.match(token)
            if token.startswith(("**", "__")) and token.endswith(("**", "__")):
                run = self._add_run(paragraph, token[2:-2])
                run.bold = True
            elif token.startswith("$$") and token.endswith("$$"):
                self._add_run(paragraph, self._readable_latex(token[2:-2]))
            elif token.startswith("$") and token.endswith("$"):
                self._add_run(paragraph, self._readable_latex(token[1:-1]))
            elif token.startswith("`") and token.endswith("`"):
                self._add_run(paragraph, token[1:-1])
            elif token.startswith(("*", "_")) and token.endswith(("*", "_")):
                run = self._add_run(paragraph, token[1:-1])
                run.italic = True
            elif link:
                self._add_run(paragraph, f"{link.group(1)}（{link.group(2)}）")
            else:
                self._add_run(paragraph, token)

    def _add_run(self, paragraph, text: str):
        run = paragraph.add_run(text)
        run.font.name = WORD_FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), WORD_FONT_NAME)
        return run

    def _readable_latex(self, formula: str) -> str:
        text = formula.strip()
        text = re.sub(r"\\frac\{([^{}]+)\}\{([^{}]+)\}", r"(\1) / \2", text)
        for source, target in LATEX_SYMBOLS.items():
            text = text.replace(source, target)
        text = re.sub(r"\\mathrm\{([^{}]+)\}", r"\1", text)
        text = re.sub(r"\\text\{([^{}]+)\}", r"\1", text)
        text = text.replace("\\", "")
        text = re.sub(r"\^\{([^{}]+)\}", lambda match: self._superscript(match.group(1)), text)
        text = re.sub(r"\^([0-9+\-=()]+)", lambda match: self._superscript(match.group(1)), text)
        text = re.sub(r"\s+", " ", text)
        return text

    def _superscript(self, value: str) -> str:
        if all(char in "0123456789+-=()" for char in value):
            return value.translate(SUPERSCRIPT_DIGITS)
        return f"^({value})"
