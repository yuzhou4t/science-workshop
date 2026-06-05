from docx import Document
from docx.oxml.ns import qn

from app.services.docx_exporter import DocxExporter


def test_export_markdown_to_docx_preserves_headings_and_paragraphs(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown("# 主标题\n\n## 小标题\n\n这是一段正文。", output)

    assert output.exists()
    doc = Document(output)
    texts = [paragraph.text for paragraph in doc.paragraphs]
    assert "主标题" in texts
    assert "小标题" in texts
    assert "这是一段正文。" in texts


def test_export_markdown_to_docx_converts_common_markdown_markers(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown(
        "# 主标题\n\n"
        "### 三级标题\n\n"
        "**核心发现**：这是正文。\n\n"
        "1. 第一条\n"
        "2. 第二条\n"
        "- 要点一\n",
        output,
    )

    doc = Document(output)
    texts = [paragraph.text for paragraph in doc.paragraphs]
    joined = "\n".join(texts)
    assert "主标题" in texts
    assert "三级标题" in texts
    assert "核心发现：这是正文。" in texts
    assert "第一条" in texts
    assert "第二条" in texts
    assert "要点一" in texts
    assert "#" not in joined
    assert "**" not in joined
    assert "1. 第一条" not in joined


def test_export_markdown_to_docx_makes_latex_formula_readable(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown(
        "核心公式：$A_B \\approx k^2 / n_b$。\n\n"
        "$$\n"
        "PRR_r = \\frac{RC_r - RU_r}{RD_r}\n"
        "$$\n",
        output,
    )

    doc = Document(output)
    joined = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "A_B ≈ k² / n_b" in joined
    assert "PRR_r = (RC_r - RU_r) / RD_r" in joined
    assert "\\frac" not in joined
    assert "$" not in joined
    assert "^2" not in joined


def test_export_markdown_to_docx_uses_simsun_for_document_styles(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown("# 主标题\n\n这是一段正文。", output)

    doc = Document(output)
    normal_fonts = doc.styles["Normal"].element.rPr.rFonts
    title_fonts = doc.styles["Title"].element.rPr.rFonts
    assert doc.styles["Normal"].font.name == "SimSun"
    assert normal_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert doc.styles["Title"].font.name == "SimSun"
    assert title_fonts.get(qn("w:eastAsia")) == "SimSun"
    assert doc.paragraphs[0].style.name == "Title"


def test_export_markdown_to_docx_inserts_default_editable_title_when_missing(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown("这是一段没有标题的正文。", output)

    doc = Document(output)
    assert doc.paragraphs[0].text == "研读非洲｜第X期"
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.core_properties.title == "研读非洲｜第X期"
    assert doc.paragraphs[1].text == "这是一段没有标题的正文。"


def test_export_markdown_to_docx_uses_public_account_title_line_as_word_title(tmp_path) -> None:
    output = tmp_path / "final.docx"
    exporter = DocxExporter()

    exporter.export_markdown("【研读非洲｜第10期】从支付普惠到金融深化\n\n图片\n\n正文。", output)

    doc = Document(output)
    assert doc.paragraphs[0].text == "【研读非洲｜第10期】从支付普惠到金融深化"
    assert doc.paragraphs[0].style.name == "Title"
    assert doc.core_properties.title == "【研读非洲｜第10期】从支付普惠到金融深化"
    assert [paragraph.text for paragraph in doc.paragraphs].count("研读非洲｜第X期") == 0
